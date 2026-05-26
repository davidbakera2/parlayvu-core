"""Plain-text extraction from common client file formats.

One shared module so both Nathan's on-demand read_client_file tool and the
client_file_ingester service use exactly the same extraction logic. If we ever
add xlsx, pptx, or OCR support, it lands here once and both callers inherit it.
"""
from __future__ import annotations

import logging
from io import BytesIO
from typing import Callable, Optional

logger = logging.getLogger("parlayvu.tools.text_extractors")


def decode_text(data: bytes) -> str:
    """Best-effort UTF-8 decode, replacing undecodable bytes."""
    return data.decode("utf-8", errors="replace")


def extract_pdf_text(data: bytes) -> str:
    """Extract concatenated text from a PDF's pages.

    Returns an empty string for image-only PDFs without embedded text. Raises
    ValueError if the bytes are not a parseable PDF.
    """
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError(
            "pypdf is required for PDF extraction. Add `pypdf>=4.0` to requirements.txt."
        ) from exc

    try:
        reader = PdfReader(BytesIO(data))
    except Exception as exc:
        raise ValueError(f"PDF could not be parsed: {exc}") from exc

    pages_text: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
        except Exception as exc:
            logger.warning("PDF page %d extraction failed: %s", i + 1, exc)
            text = ""
        if text.strip():
            pages_text.append(f"--- page {i + 1} ---\n{text.strip()}")
    return "\n\n".join(pages_text)


def extract_docx_text(data: bytes) -> str:
    """Extract concatenated text + table cells from a .docx via python-docx.

    Raises ValueError if the bytes are not a parseable .docx.
    """
    from docx import Document  # python-docx is already in requirements.txt

    try:
        doc = Document(BytesIO(data))
    except Exception as exc:
        raise ValueError(f"DOCX could not be parsed: {exc}") from exc

    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            parts.append(text)
    # Tables aren't in doc.paragraphs — collect them too.
    for table in doc.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            row_text = " | ".join(c for c in cells if c)
            if row_text:
                parts.append(row_text)
    return "\n\n".join(parts)


Extractor = Callable[[bytes], str]


def detect_extractor(path: str) -> Optional[Extractor]:
    """Return the right text extractor for the file extension, or None for
    types we don't currently support."""
    lower = path.lower()
    if lower.endswith((".md", ".markdown", ".txt", ".rst")):
        return decode_text
    if lower.endswith(".pdf"):
        return extract_pdf_text
    if lower.endswith(".docx"):
        return extract_docx_text
    return None
