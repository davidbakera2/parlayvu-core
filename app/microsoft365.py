import os
import re
import zipfile
from dataclasses import dataclass
from datetime import datetime, timezone
from html import escape
from io import BytesIO
from typing import Any, Optional
from urllib.parse import quote

import httpx

AGENT_MAILBOX_ENV = {
    "nathan": "NATHAN_MAILBOX",
    "alex": "ALEX_MAILBOX",
    "ava": "AVA_MAILBOX",
    "blake": "BLAKE_MAILBOX",
    "casey": "CASEY_MAILBOX",
    "codey": "CODEY_MAILBOX",
    "dylan": "DYLAN_MAILBOX",
    "jordan": "JORDAN_MAILBOX",
    "michael": "MICHAEL_MAILBOX",
    "morgan": "MORGAN_MAILBOX",
    "nora": "NORA_MAILBOX",
    "riley": "RILEY_MAILBOX",
    "taylor": "TAYLOR_MAILBOX",
}


@dataclass(frozen=True)
class Microsoft365Settings:
    tenant_id: str
    client_id: str
    client_secret: str
    graph_scope: str
    webhook_client_state: str
    allow_send: bool
    agent_mailboxes: dict[str, str]
    onenote_owner_mailbox: str = ""
    onenote_notebook_id: str = ""
    onenote_notebook_name: str = "RamAir"
    onenote_section_id: str = ""
    onenote_section_name: str = "Meeting Notes"
    files_drive_id: str = ""
    files_folder_item_id: str = ""
    files_base_path: str = "03_Deliverables/Meeting Notes"

    @property
    def configured(self) -> bool:
        return bool(self.tenant_id and self.client_id and self.client_secret)


def _truthy(value: str) -> bool:
    return value.strip().lower() in {"1", "true", "yes", "on"}


def get_microsoft365_settings() -> Microsoft365Settings:
    agent_mailboxes = {
        agent: mailbox
        for agent, env_name in AGENT_MAILBOX_ENV.items()
        if (mailbox := os.getenv(env_name, "").strip())
    }
    return Microsoft365Settings(
        tenant_id=os.getenv("MICROSOFT_TENANT_ID", ""),
        client_id=os.getenv("MICROSOFT_CLIENT_ID", ""),
        client_secret=os.getenv("MICROSOFT_CLIENT_SECRET", ""),
        graph_scope=os.getenv("MICROSOFT_GRAPH_SCOPE", "https://graph.microsoft.com/.default"),
        webhook_client_state=os.getenv("MICROSOFT_WEBHOOK_CLIENT_STATE", ""),
        allow_send=_truthy(os.getenv("MICROSOFT_GRAPH_ALLOW_SEND", "false")),
        agent_mailboxes=agent_mailboxes,
        onenote_owner_mailbox=os.getenv("ONENOTE_OWNER_MAILBOX", agent_mailboxes.get("nathan", "")),
        onenote_notebook_id=os.getenv("ONENOTE_NOTEBOOK_ID", ""),
        onenote_notebook_name=os.getenv("ONENOTE_NOTEBOOK_NAME", "RamAir"),
        onenote_section_id=os.getenv("ONENOTE_SECTION_ID", ""),
        onenote_section_name=os.getenv("ONENOTE_SECTION_NAME", "Meeting Notes"),
        files_drive_id=os.getenv("M365_FILES_DRIVE_ID", ""),
        files_folder_item_id=os.getenv("M365_FILES_FOLDER_ITEM_ID", ""),
        files_base_path=os.getenv("M365_FILES_BASE_PATH", "03_Deliverables/Meeting Notes"),
    )


def mailbox_status(settings: Optional[Microsoft365Settings] = None) -> dict[str, Any]:
    active_settings = settings or get_microsoft365_settings()
    return {
        "configured": active_settings.configured,
        "allow_send": active_settings.allow_send,
        "onenote": {
            "configured": active_settings.configured
            and bool(active_settings.onenote_owner_mailbox)
            and bool(active_settings.onenote_section_id or active_settings.onenote_section_name),
            "owner_mailbox": active_settings.onenote_owner_mailbox or None,
            "notebook_id_configured": bool(active_settings.onenote_notebook_id),
            "notebook_name": active_settings.onenote_notebook_name or None,
            "section_id_configured": bool(active_settings.onenote_section_id),
            "section_name": active_settings.onenote_section_name or None,
        },
        "files": {
            "configured": active_settings.configured and bool(active_settings.files_drive_id),
            "drive_id_configured": bool(active_settings.files_drive_id),
            "folder_item_id_configured": bool(active_settings.files_folder_item_id),
            "base_path": active_settings.files_base_path or None,
        },
        "agents": {
            agent: {
                "configured": agent in active_settings.agent_mailboxes,
                "mailbox": active_settings.agent_mailboxes.get(agent),
            }
            for agent in AGENT_MAILBOX_ENV
        },
    }


def sanitize_file_stem(value: str, *, default: str = "meeting-notes") -> str:
    stem = re.sub(r"[^A-Za-z0-9._ -]+", "", value.strip()).strip(" ._-")
    stem = re.sub(r"\s+", "-", stem).lower()
    return stem[:80] or default


def build_meeting_notes_markdown(
    *,
    title: str,
    summary: str,
    client_id: Optional[str] = None,
    project_id: Optional[str] = None,
    source_label: str = "ParlayVU project memory",
) -> str:
    now = datetime.now(timezone.utc).replace(microsecond=0).isoformat()
    lines = [
        f"# {title.strip() or 'Meeting Notes'}",
        "",
        f"- Source of truth: {source_label}",
        "- Published by: ParlayVU Nathan",
        f"- Published at: {now}",
    ]
    if client_id:
        lines.append(f"- Client: {client_id}")
    if project_id:
        lines.append(f"- Project: {project_id}")
    lines.extend(["", "## Meeting Summary", "", summary.strip() or "No summary provided.", ""])
    return "\n".join(lines)


def build_meeting_notes_docx(
    *,
    title: str,
    summary: str,
    client_id: Optional[str] = None,
    project_id: Optional[str] = None,
    source_label: str = "ParlayVU project memory",
) -> bytes:
    metadata = [
        f"Source of truth: {source_label}",
        "Published by: ParlayVU Nathan",
    ]
    if client_id:
        metadata.append(f"Client: {client_id}")
    if project_id:
        metadata.append(f"Project: {project_id}")

    def paragraph(text: str, style: str | None = None) -> str:
        style_xml = f'<w:pPr><w:pStyle w:val="{style}"/></w:pPr>' if style else ""
        return f"<w:p>{style_xml}<w:r><w:t>{escape(text)}</w:t></w:r></w:p>"

    body = [paragraph(title.strip() or "Meeting Notes", "Title")]
    body.extend(paragraph(item) for item in metadata)
    body.append(paragraph("Meeting Summary", "Heading1"))
    body.extend(paragraph(part) for part in (summary.strip() or "No summary provided.").splitlines() if part.strip())

    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f"<w:body>{''.join(body)}<w:sectPr/></w:body>"
        "</w:document>"
    )
    content_types_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/word/document.xml" '
        'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
        "</Types>"
    )
    rels_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
        'Target="word/document.xml"/>'
        "</Relationships>"
    )

    buffer = BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", content_types_xml)
        archive.writestr("_rels/.rels", rels_xml)
        archive.writestr("word/document.xml", document_xml)
    return buffer.getvalue()


def build_strategy_docx(
    *,
    meeting_title: str,
    blake_analysis: dict,
    nathan_strategy: str,
    client_id: Optional[str] = None,
    project_id: Optional[str] = None,
) -> bytes:
    """Build a rich strategy .docx from Blake's analysis and Nathan's strategy text."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    doc = Document()
    now = datetime.now(timezone.utc).strftime("%B %d, %Y at %H:%M UTC")

    # --- Title ---
    doc.add_heading(f"{meeting_title} — Strategy & Next Steps", 0)

    # --- Metadata ---
    meta = doc.add_paragraph()
    meta.add_run(f"Date: {now}\n")
    if client_id:
        meta.add_run(f"Client: {client_id}\n")
    if project_id:
        meta.add_run(f"Project: {project_id}\n")
    meta.add_run("Prepared by: Nathan Ellis, ParlayVU.ai")

    doc.add_paragraph()  # spacer

    # --- Meeting Summary ---
    summary = blake_analysis.get("meeting_summary", "").strip()
    if summary:
        doc.add_heading("Meeting Summary", 1)
        for para in summary.splitlines():
            if para.strip():
                doc.add_paragraph(para.strip())

    # --- Key Themes ---
    themes = blake_analysis.get("key_themes", [])
    if themes:
        doc.add_heading("Key Themes", 1)
        for theme in themes:
            doc.add_paragraph(str(theme).strip(), style="List Bullet")

    # --- Decisions Made ---
    decisions = blake_analysis.get("decisions_made", [])
    if decisions:
        doc.add_heading("Decisions Made", 1)
        for d in decisions:
            doc.add_paragraph(str(d).strip(), style="List Bullet")

    # --- Action Items table ---
    action_items = blake_analysis.get("action_items", [])
    if action_items:
        doc.add_heading("Action Items", 1)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        for cell, label in zip(hdr_cells, ["Action Item", "Owner", "Priority"]):
            cell.text = label
            run = cell.paragraphs[0].runs[0]
            run.bold = True
        for item in action_items:
            row = table.add_row().cells
            row[0].text = str(item.get("item", ""))
            row[1].text = str(item.get("owner", ""))
            row[2].text = str(item.get("priority", "")).capitalize()

    # --- Open Questions ---
    questions = blake_analysis.get("open_questions", [])
    if questions:
        doc.add_heading("Open Questions", 1)
        for q in questions:
            doc.add_paragraph(str(q).strip(), style="List Bullet")

    # --- Strategic Opportunities ---
    opportunities = blake_analysis.get("strategic_opportunities", [])
    if opportunities:
        doc.add_heading("Strategic Opportunities", 1)
        for o in opportunities:
            doc.add_paragraph(str(o).strip(), style="List Bullet")

    # --- Nathan's Strategy (parse markdown sections) ---
    if nathan_strategy:
        doc.add_heading("Nathan's Strategic Assessment", 1)
        _render_strategy_markdown(doc, nathan_strategy)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _render_strategy_markdown(doc, text: str) -> None:
    """Render Nathan's markdown strategy text into Word paragraphs."""
    pending: list[str] = []

    def flush():
        for line in pending:
            if line.strip():
                doc.add_paragraph(line.strip())
        pending.clear()

    for line in text.splitlines():
        if line.startswith("## "):
            flush()
            doc.add_heading(line[3:].strip(), 2)
        elif line.startswith("### "):
            flush()
            doc.add_heading(line[4:].strip(), 3)
        elif line.startswith(("- ", "* ")):
            flush()
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line and line[0].isdigit() and ". " in line[:5]:
            flush()
            doc.add_paragraph(line.split(". ", 1)[1].strip(), style="List Number")
        else:
            pending.append(line)

    flush()


def _summary_sections(summary: str) -> dict[str, str]:
    sections: dict[str, list[str]] = {}
    current = "summary"
    for raw_line in summary.splitlines():
        line = raw_line.strip()
        heading = re.match(r"^(?:#{1,6}\s*)?([A-Za-z][A-Za-z /_-]{1,40}):?\s*$", line)
        if heading:
            normalized = re.sub(r"[^a-z0-9]+", "_", heading.group(1).strip().lower()).strip("_")
            if normalized in {
                "attendees",
                "summary",
                "meeting_summary",
                "decisions",
                "action_items",
                "actions",
                "questions",
                "open_questions",
                "next_steps",
            }:
                current = normalized
                sections.setdefault(current, [])
                continue
        sections.setdefault(current, []).append(raw_line)

    return {
        key: "\n".join(part for part in parts).strip()
        for key, parts in sections.items()
        if "\n".join(parts).strip()
    }


def _default_meeting_date_time() -> str:
    """Friendly date+time string used when caller doesn't supply MEETING_DATE.

    Format: "May 25, 2026 at 14:32 UTC". Uses portable strftime tokens
    (no platform-specific %-d / %#d) so it works on Windows and Unix alike,
    even though it leaves a leading zero on single-digit days.
    """
    return datetime.now(timezone.utc).strftime("%B %d, %Y at %H:%M UTC")


def build_meeting_notes_template_placeholders(
    *,
    title: str,
    summary: str,
    client_id: Optional[str] = None,
    client_name: Optional[str] = None,
    client_full_name: Optional[str] = None,
    project_id: Optional[str] = None,
    project_name: Optional[str] = None,
    meeting_date_time: Optional[str] = None,
    source_label: str = "ParlayVU project memory",
) -> dict[str, str]:
    """Build the simple-text placeholder map for the meeting notes template.

    This returns only the single-value substitutions. The structured
    fields (attendees, decisions, action items, etc.) are passed to the
    renderer via separate `list_items` and `action_items` arguments, see
    `render_meeting_notes_template_docx`.

    Backward compatibility: callers that don't supply the new structured
    fields will see the same legacy keys (`{{ATTENDEES}}`, `{{DECISIONS}}`,
    etc.) populated from the markdown summary as before. New callers can
    omit those — the renderer will handle list duplication instead.
    """
    sections = _summary_sections(summary)
    client = (client_name or client_id or "RamAir").strip() or "RamAir"
    client_full = (client_full_name or client).strip() or client
    project = project_name or project_id or client_id or "RamAir"
    meeting_date = meeting_date_time or _default_meeting_date_time()

    def first_section(*keys: str, default: str = "Not specified.") -> str:
        for key in keys:
            if sections.get(key):
                return sections[key]
        return default

    return {
        "{{MEETING_TITLE}}": title.strip() or "Meeting Notes",
        "{{MEETING_DATE}}": meeting_date,
        "{{CLIENT}}": client,
        "{{CLIENT_NAME}}": client_full,
        "{{PROJECT}}": project,
        # Legacy single-value placeholders. New callers that pass
        # structured `list_items` to the renderer can leave these empty -
        # the renderer's list-duplication path takes precedence.
        "{{ATTENDEES}}": first_section("attendees"),
        "{{SUMMARY}}": first_section("meeting_summary", "summary", default=summary.strip() or "No summary provided."),
        "{{DECISIONS}}": first_section("decisions"),
        "{{ACTION_ITEMS}}": first_section("action_items", "actions"),
        "{{QUESTIONS}}": first_section("questions", "open_questions"),
        "{{NEXT_STEPS}}": first_section("next_steps"),
        "{{SOURCE_MATERIAL}}": source_label,
    }


def _replace_docx_placeholder(xml: str, placeholder: str, value: str) -> tuple[str, int]:
    """Legacy XML-string substitution. Used as a fallback for the headers/footers
    path in render_meeting_notes_template_docx for any placeholders not handled
    by the structured renderer."""
    escaped_value = escape(value)
    direct_count = xml.count(placeholder)
    updated = xml.replace(placeholder, escaped_value)

    xml_tag_gap = r"(?:<[^>]+>)*"
    split_pattern = xml_tag_gap.join(re.escape(char) for char in placeholder)
    updated, split_count = re.subn(split_pattern, lambda _match: escaped_value, updated)
    return updated, direct_count + split_count


# ── Structured DOCX template rendering ────────────────────────────────────────
# Supports three behaviors:
#   1. Simple text substitution: {{KEY}} → "value"
#   2. List paragraph duplication: a paragraph containing {{KEY}} is duplicated
#      once per item in list_items[{{KEY}}]. Preserves paragraph style, so
#      bulleted-list paragraphs in the template stay bulleted in the output.
#   3. Action items table row duplication: a table row containing any of
#      {{ACTION_OWNER}}, {{ACTION_ITEM}}, {{ACTION_DUE}} is duplicated once
#      per action item, with the three placeholders filled in each copy.
# All three behaviors work in body paragraphs, table cells, and headers/footers.

# Each column accepts a few natural-feeling placeholder names so the template
# author isn't tied to one exact spelling. The renderer substitutes ALL
# variants in each row, so a template using {{ACTION_DATE}} works the same
# as one using {{ACTION_DUE}} or {{DUE_DATE}}.
_ACTION_OWNER_PLACEHOLDERS = ("{{ACTION_OWNER}}", "{{OWNER}}")
_ACTION_ITEM_PLACEHOLDERS  = ("{{ACTION_ITEM}}", "{{ACTION}}", "{{ITEM}}", "{{TASK}}")
_ACTION_DUE_PLACEHOLDERS   = ("{{ACTION_DUE}}", "{{ACTION_DATE}}", "{{DUE_DATE}}", "{{DUE}}")
_ALL_ACTION_PLACEHOLDERS = (
    _ACTION_OWNER_PLACEHOLDERS + _ACTION_ITEM_PLACEHOLDERS + _ACTION_DUE_PLACEHOLDERS
)


def _substitute_in_paragraph(paragraph, placeholder: str, value: str) -> int:
    """Replace placeholder text in a paragraph, handling cases where the
    placeholder spans multiple runs. Preserves paragraph-level styling
    (list bullets, indent) but collapses run-level formatting *inside* the
    placeholder to whatever the first run had."""
    full_text = "".join(r.text or "" for r in paragraph.runs)
    if placeholder not in full_text:
        return 0
    count = full_text.count(placeholder)
    new_text = full_text.replace(placeholder, value)
    runs = list(paragraph.runs)
    if runs:
        runs[0].text = new_text
        for r in runs[1:]:
            r.text = ""
    return count


def _all_paragraphs(doc):
    """Yield every paragraph anywhere in the document: body, table cells,
    headers, footers, and nested tables-within-cells."""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
                for nested in cell.tables:
                    for nrow in nested.rows:
                        for ncell in nrow.cells:
                            for np in ncell.paragraphs:
                                yield np
    for section in doc.sections:
        for hf in (section.header, section.footer):
            for p in hf.paragraphs:
                yield p
            for table in hf.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            yield p


def _substitute_all(doc, placeholder: str, value: str) -> int:
    count = 0
    for p in _all_paragraphs(doc):
        count += _substitute_in_paragraph(p, placeholder, value)
    return count


def _substitute_text_in_element(element, placeholder: str, value: str) -> int:
    """Replace placeholder in all <w:t> text nodes within an XML element.
    Used for row-duplication where we operate at the XML level after deepcopy."""
    from docx.oxml.ns import qn  # local import to avoid hard dep at module load
    count = 0
    for t in element.iter(qn("w:t")):
        if t.text and placeholder in t.text:
            count += t.text.count(placeholder)
            t.text = t.text.replace(placeholder, value)
    return count


def _duplicate_list_paragraphs_in_doc(doc, placeholder: str, items: list[str]) -> int:
    """Find every paragraph containing the placeholder; duplicate it once per
    item (preserving paragraph style); fill each copy with one item. If
    items is empty, remove the placeholder paragraph entirely.

    Operates on body, table cells, headers, footers.
    """
    from copy import deepcopy
    from docx.text.paragraph import Paragraph

    count = 0

    # Snapshot paragraphs first so we don't iterate over a mutating collection
    targets = []
    for p in _all_paragraphs(doc):
        if placeholder in (p.text or ""):
            targets.append(p)

    for p in targets:
        element = p._element
        parent_element = element.getparent()
        if parent_element is None:
            continue

        if not items:
            parent_element.remove(element)
            count += 1
            continue

        # Snapshot the placeholder-bearing element BEFORE substituting, so
        # each copy still has the raw placeholder we can target.
        template_element = deepcopy(element)

        # First item goes into the original paragraph (preserves any
        # surrounding context / order)
        _substitute_in_paragraph(p, placeholder, str(items[0]))
        count += 1

        # Additional items: deep-copy the SNAPSHOT (still has placeholder)
        # and insert each after the original
        idx = list(parent_element).index(element)
        for offset, item in enumerate(items[1:], start=1):
            new_element = deepcopy(template_element)
            parent_element.insert(idx + offset, new_element)
            new_p = Paragraph(new_element, p._parent)
            _substitute_in_paragraph(new_p, placeholder, str(item))
            count += 1

    return count


def _action_item_owner(item: dict[str, str]) -> str:
    return str(item.get("owner") or item.get("assignee") or "")


def _action_item_action(item: dict[str, str]) -> str:
    return str(item.get("action") or item.get("item") or item.get("task") or "")


def _action_item_due(item: dict[str, str]) -> str:
    return str(item.get("due_date") or item.get("due") or "")


def _duplicate_action_item_rows(doc, action_items: list[dict[str, str]]) -> int:
    """Find table rows containing any of the action-item placeholders and
    duplicate them once per action item. If action_items is empty, fills
    the row with em-dashes so the table doesn't show literal placeholders.

    The renderer accepts multiple natural placeholder names per column
    (see _ACTION_OWNER_PLACEHOLDERS etc.) so the template author can use
    {{ACTION_DATE}}, {{DUE_DATE}}, or {{ACTION_DUE}} interchangeably.
    """
    from copy import deepcopy

    count = 0

    def row_contains_placeholders(row) -> bool:
        text = " ".join(cell.text or "" for cell in row.cells)
        return any(ph in text for ph in _ALL_ACTION_PLACEHOLDERS)

    for table in doc.tables:
        template_row = None
        for row in table.rows:
            if row_contains_placeholders(row):
                template_row = row
                break
        if template_row is None:
            continue

        tr_element = template_row._element
        tbl_element = tr_element.getparent()
        if tbl_element is None:
            continue

        def fill_row(row, owner: str, action: str, due: str) -> None:
            """Substitute every accepted placeholder variant in the row with the
            corresponding value. Variants we don't find are simply no-ops."""
            for cell in row.cells:
                for p in cell.paragraphs:
                    for ph in _ACTION_OWNER_PLACEHOLDERS:
                        _substitute_in_paragraph(p, ph, owner)
                    for ph in _ACTION_ITEM_PLACEHOLDERS:
                        _substitute_in_paragraph(p, ph, action)
                    for ph in _ACTION_DUE_PLACEHOLDERS:
                        _substitute_in_paragraph(p, ph, due)

        if not action_items:
            fill_row(template_row, "—", "No action items recorded.", "—")
            count += 1
            continue

        # Snapshot the template row BEFORE mutating the original. Each
        # subsequent row is a deep-copy of this snapshot, so all rows
        # start with the same placeholder text and get filled independently.
        from docx.table import _Row  # type: ignore
        template_snapshot = deepcopy(tr_element)

        # First item: fill the original row in place
        first = action_items[0]
        fill_row(template_row, _action_item_owner(first), _action_item_action(first), _action_item_due(first))
        count += 1

        # Remaining items: clone the snapshot and fill
        template_idx = list(tbl_element).index(tr_element)
        for offset, item in enumerate(action_items[1:], start=1):
            new_tr = deepcopy(template_snapshot)
            tbl_element.insert(template_idx + offset, new_tr)
            # Wrap the inserted <w:tr> as a python-docx Row so fill_row can
            # iterate its cells. _Row takes (tr_element, parent_table).
            new_row = _Row(new_tr, table)
            fill_row(new_row, _action_item_owner(item), _action_item_action(item), _action_item_due(item))
            count += 1

    return count


def render_meeting_notes_template_docx(
    template_docx: bytes,
    placeholders: dict[str, str],
    *,
    list_items: Optional[dict[str, list[str]]] = None,
    action_items: Optional[list[dict[str, str]]] = None,
) -> bytes:
    """Render the meeting notes DOCX template with substitution + duplication.

    Two-phase rendering:

      Phase 1 - structural transforms via python-docx (only run if list_items
      or action_items are provided). Handles:
        * List paragraph duplication: a paragraph containing a list-item
          placeholder is duplicated once per item, preserving paragraph
          style (bulleted-list paragraphs stay bulleted).
        * Action-items table row duplication: the first row containing
          {{ACTION_OWNER}}/{{ACTION_ITEM}}/{{ACTION_DUE}} is duplicated
          per action item.

      Phase 2 - simple text substitution via XML string replace. Iterates
      every word/*.xml file in the zip (body, headers, footers, footnotes,
      etc.) and does literal placeholder->value replacement. This catches
      everything Phase 1 didn't, including headers and footers that aren't
      reachable through python-docx's section API.

    Args:
        template_docx: The .docx template file contents as bytes.
        placeholders: Simple text substitutions: {"{{KEY}}": "value"}.
        list_items: Optional dict of {"{{KEY}}": [items...]} for paragraph
            duplication. Empty list deletes the placeholder paragraph.
        action_items: Optional list of {owner, action, due_date} dicts for
            table row duplication.

    Returns:
        The rendered .docx file as bytes.
    """
    if not template_docx:
        raise ValueError("Meeting notes template is empty")

    structural_count = 0

    # ---- Phase 1: python-docx structural transforms ----
    if list_items or action_items is not None:
        from docx import Document  # lazy import
        try:
            doc = Document(BytesIO(template_docx))
        except Exception as exc:
            raise ValueError("Meeting notes template is not a valid DOCX file") from exc

        if list_items:
            for placeholder, items in list_items.items():
                structural_count += _duplicate_list_paragraphs_in_doc(doc, placeholder, list(items or []))

        if action_items is not None:
            structural_count += _duplicate_action_item_rows(doc, list(action_items))

        intermediate = BytesIO()
        doc.save(intermediate)
        template_docx = intermediate.getvalue()

    # ---- Phase 2: simple XML-level string substitution ----
    output = BytesIO()
    text_count = 0
    try:
        with zipfile.ZipFile(BytesIO(template_docx), "r") as source:
            names = set(source.namelist())
            if "word/document.xml" not in names:
                raise ValueError("Meeting notes template is missing word/document.xml")

            with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as target:
                for info in source.infolist():
                    content = source.read(info.filename)
                    if info.filename.startswith("word/") and info.filename.endswith(".xml"):
                        text = content.decode("utf-8")
                        for placeholder, value in (placeholders or {}).items():
                            text, count = _replace_docx_placeholder(text, placeholder, value or "")
                            text_count += count
                        content = text.encode("utf-8")
                    target.writestr(info, content)
    except zipfile.BadZipFile as exc:
        raise ValueError("Meeting notes template is not a valid DOCX zip") from exc
    except UnicodeDecodeError as exc:
        raise ValueError("Meeting notes template contains unsupported XML encoding") from exc

    if structural_count == 0 and text_count == 0:
        raise ValueError("Meeting notes template did not contain recognized placeholders")
    return output.getvalue()


def build_onenote_meeting_page_html(
    *,
    title: str,
    summary: str,
    client_id: Optional[str] = None,
    project_id: Optional[str] = None,
    source_label: str = "ParlayVU project memory",
) -> str:
    safe_title = escape(title.strip() or "Meeting Notes")
    paragraphs = [
        f"<p>{escape(part).replace(chr(10), '<br />')}</p>"
        for part in summary.strip().split("\n\n")
        if part.strip()
    ]
    metadata_items = [
        ("Source of truth", source_label),
        ("Client", client_id),
        ("Project", project_id),
    ]
    metadata = "\n".join(
        f"<li><strong>{escape(label)}:</strong> {escape(value)}</li>"
        for label, value in metadata_items
        if value
    )
    return (
        "<!DOCTYPE html>\n"
        "<html>\n"
        "<head>\n"
        f"  <title>{safe_title}</title>\n"
        '  <meta name="created-by" content="ParlayVU Nathan" />\n'
        "</head>\n"
        "<body>\n"
        f"  <h1>{safe_title}</h1>\n"
        f"  <ul>{metadata}</ul>\n"
        "  <h2>Meeting Summary</h2>\n"
        f"  {''.join(paragraphs) or '<p>No summary provided.</p>'}\n"
        "</body>\n"
        "</html>"
    )


class MicrosoftGraphClient:
    def __init__(self, settings: Optional[Microsoft365Settings] = None):
        self.settings = settings or get_microsoft365_settings()

    async def get_access_token(self) -> str:
        if not self.settings.configured:
            raise RuntimeError("Microsoft 365 Graph credentials are not configured")

        token_url = f"https://login.microsoftonline.com/{self.settings.tenant_id}/oauth2/v2.0/token"
        data = {
            "client_id": self.settings.client_id,
            "client_secret": self.settings.client_secret,
            "scope": self.settings.graph_scope,
            "grant_type": "client_credentials",
        }
        async with httpx.AsyncClient(timeout=20) as client:
            response = await client.post(token_url, data=data)
            response.raise_for_status()
            payload = response.json()
        return payload["access_token"]

    def mailbox_for_agent(self, agent_name: str) -> str:
        mailbox = self.settings.agent_mailboxes.get(agent_name.lower())
        if not mailbox:
            raise ValueError(f"No Microsoft 365 mailbox is configured for agent: {agent_name}")
        return mailbox

    def onenote_owner_mailbox(self) -> str:
        mailbox = self.settings.onenote_owner_mailbox or self.settings.agent_mailboxes.get("nathan")
        if not mailbox:
            raise ValueError("No OneNote owner mailbox is configured")
        return mailbox

    async def _graph_get(self, path: str, token: str) -> dict[str, Any]:
        async with httpx.AsyncClient(timeout=20) as client:
            response = await client.get(
                f"https://graph.microsoft.com/v1.0{path}",
                headers={"Authorization": f"Bearer {token}"},
            )
            response.raise_for_status()
            return response.json()

    async def resolve_channel_files_folder(
        self,
        *,
        team_id: Optional[str] = None,
        channel_id: Optional[str] = None,
        token: Optional[str] = None,
    ) -> dict[str, Any]:
        resolved_team_id = team_id
        resolved_channel_id = channel_id
        if not resolved_team_id or not resolved_channel_id:
            raise ValueError("Teams file publishing requires a team_id and channel_id")

        access_token = token or await self.get_access_token()
        folder = await self._graph_get(
            f"/teams/{quote(resolved_team_id, safe='')}/channels/{quote(resolved_channel_id, safe='')}/filesFolder",
            access_token,
        )
        parent_reference = folder.get("parentReference") or {}
        drive_id = parent_reference.get("driveId")
        if not drive_id:
            raise ValueError("Graph filesFolder response did not include parentReference.driveId")
        return {
            "id": folder.get("id"),
            "name": folder.get("name"),
            "webUrl": folder.get("webUrl"),
            "drive_id": drive_id,
            "team_id": resolved_team_id,
            "channel_id": resolved_channel_id,
        }

    async def list_channel_files(
        self,
        *,
        team_id: str,
        channel_id: str,
        folder_path: Optional[str] = None,
        token: Optional[str] = None,
    ) -> list[dict[str, Any]]:
        """List the files and subfolders at a path inside a Teams channel's Files tab.

        `folder_path` is relative to the channel's Files root (e.g. "Reports" or
        "03_Deliverables/Meeting Notes"). When None or empty, lists the root.
        Returns a list of dicts shaped like {name, kind, size, last_modified,
        web_url, path}; `kind` is "file" or "folder" so callers can recurse.
        """
        access_token = token or await self.get_access_token()
        folder = await self.resolve_channel_files_folder(
            team_id=team_id, channel_id=channel_id, token=access_token
        )
        drive_id = folder["drive_id"]
        root_item_id = folder["id"]

        clean_path = (folder_path or "").strip("/")
        if clean_path:
            quoted_path = quote(clean_path, safe="/")
            list_path = f"/drives/{quote(drive_id, safe='')}/root:/{quoted_path}:/children"
        else:
            list_path = f"/drives/{quote(drive_id, safe='')}/items/{quote(root_item_id, safe='')}/children"

        response = await self._graph_get(
            f"{list_path}?$top=200&$select=id,name,size,lastModifiedDateTime,file,folder,webUrl",
            access_token,
        )
        items: list[dict[str, Any]] = []
        for entry in response.get("value", []):
            kind = "folder" if entry.get("folder") else "file" if entry.get("file") else "other"
            name = entry.get("name") or ""
            child_path = f"{clean_path}/{name}".strip("/") if clean_path else name
            items.append(
                {
                    "name": name,
                    "kind": kind,
                    "size": entry.get("size"),
                    "last_modified": entry.get("lastModifiedDateTime"),
                    "web_url": entry.get("webUrl"),
                    "path": child_path,
                }
            )
        return items

    async def upload_drive_file(
        self,
        *,
        drive_id: str,
        filename: str,
        content: bytes,
        folder_path: Optional[str] = None,
        folder_item_id: Optional[str] = None,
        content_type: str = "application/octet-stream",
    ) -> dict[str, Any]:
        if not drive_id:
            raise ValueError("drive_id is required for SharePoint file upload")
        clean_filename = filename.strip()
        if not clean_filename:
            raise ValueError("filename is required for SharePoint file upload")

        token = await self.get_access_token()
        relative_path = "/".join(part.strip("/") for part in [folder_path or "", clean_filename] if part.strip("/"))
        quoted_path = quote(relative_path, safe="/")
        if folder_item_id:
            path = f"/drives/{quote(drive_id, safe='')}/items/{quote(folder_item_id, safe='')}:/{quoted_path}:/content"
        else:
            path = f"/drives/{quote(drive_id, safe='')}/root:/{quoted_path}:/content"

        async with httpx.AsyncClient(timeout=30) as client:
            response = await client.put(
                f"https://graph.microsoft.com/v1.0{path}",
                headers={
                    "Authorization": f"Bearer {token}",
                    "Content-Type": content_type,
                },
                content=content,
            )
            response.raise_for_status()
            item = response.json()
        parent_reference = item.get("parentReference") or {}
        return {
            "id": item.get("id"),
            "name": item.get("name") or clean_filename,
            "webUrl": item.get("webUrl"),
            "size": item.get("size"),
            "drive_id": parent_reference.get("driveId") or drive_id,
            "path": relative_path,
            "content_type": content_type,
        }

    async def download_drive_file(
        self,
        *,
        drive_id: str,
        file_path: str,
        folder_item_id: Optional[str] = None,
    ) -> bytes:
        if not drive_id:
            raise ValueError("drive_id is required for SharePoint file download")
        clean_path = file_path.strip("/")
        if not clean_path:
            raise ValueError("file_path is required for SharePoint file download")

        token = await self.get_access_token()
        quoted_path = quote(clean_path, safe="/")
        if folder_item_id:
            path = f"/drives/{quote(drive_id, safe='')}/items/{quote(folder_item_id, safe='')}:/{quoted_path}:/content"
        else:
            path = f"/drives/{quote(drive_id, safe='')}/root:/{quoted_path}:/content"

        async with httpx.AsyncClient(timeout=30, follow_redirects=True) as client:
            response = await client.get(
                f"https://graph.microsoft.com/v1.0{path}",
                headers={"Authorization": f"Bearer {token}"},
            )
            response.raise_for_status()
            return response.content

    async def upload_teams_channel_file(
        self,
        *,
        filename: str,
        content: bytes,
        content_type: str,
        team_id: Optional[str] = None,
        channel_id: Optional[str] = None,
        folder_path: Optional[str] = None,
    ) -> dict[str, Any]:
        if team_id and channel_id:
            folder = await self.resolve_channel_files_folder(team_id=team_id, channel_id=channel_id)
            upload = await self.upload_drive_file(
                drive_id=folder["drive_id"],
                folder_item_id=folder["id"],
                folder_path=folder_path or self.settings.files_base_path,
                filename=filename,
                content=content,
                content_type=content_type,
            )
            return {**upload, "team_id": folder["team_id"], "channel_id": folder["channel_id"], "files_folder": folder}

        if self.settings.files_drive_id:
            return await self.upload_drive_file(
                drive_id=self.settings.files_drive_id,
                folder_item_id=self.settings.files_folder_item_id or None,
                folder_path=folder_path or self.settings.files_base_path,
                filename=filename,
                content=content,
                content_type=content_type,
            )

        raise ValueError("Teams/SharePoint file publishing is not configured")

    async def download_teams_channel_file(
        self,
        *,
        file_path: str,
        team_id: Optional[str] = None,
        channel_id: Optional[str] = None,
    ) -> bytes:
        if team_id and channel_id:
            folder = await self.resolve_channel_files_folder(team_id=team_id, channel_id=channel_id)
            return await self.download_drive_file(
                drive_id=folder["drive_id"],
                folder_item_id=folder["id"],
                file_path=file_path,
            )

        if self.settings.files_drive_id:
            return await self.download_drive_file(
                drive_id=self.settings.files_drive_id,
                folder_item_id=self.settings.files_folder_item_id or None,
                file_path=file_path,
            )

        raise ValueError("Teams/SharePoint file publishing is not configured")

    async def resolve_onenote_section_id(self, *, token: Optional[str] = None) -> str:
        if self.settings.onenote_section_id:
            return self.settings.onenote_section_id

        mailbox = self.onenote_owner_mailbox()
        access_token = token or await self.get_access_token()
        notebook_id = self.settings.onenote_notebook_id

        if not notebook_id:
            notebooks = await self._graph_get(
                f"/users/{mailbox}/onenote/notebooks?$select=id,displayName",
                access_token,
            )
            notebook_name = self.settings.onenote_notebook_name.strip().lower()
            for notebook in notebooks.get("value", []):
                if (notebook.get("displayName") or "").strip().lower() == notebook_name:
                    notebook_id = notebook["id"]
                    break
            if not notebook_id:
                raise ValueError(f"OneNote notebook not found: {self.settings.onenote_notebook_name}")

        sections = await self._graph_get(
            f"/users/{mailbox}/onenote/notebooks/{notebook_id}/sections?$select=id,displayName",
            access_token,
        )
        section_name = self.settings.onenote_section_name.strip().lower()
        for section in sections.get("value", []):
            if (section.get("displayName") or "").strip().lower() == section_name:
                return section["id"]
        raise ValueError(f"OneNote section not found: {self.settings.onenote_section_name}")

    async def create_onenote_page(
        self,
        *,
        title: str,
        html: str,
    ) -> dict[str, Any]:
        mailbox = self.onenote_owner_mailbox()
        token = await self.get_access_token()
        section_id = await self.resolve_onenote_section_id(token=token)

        async with httpx.AsyncClient(timeout=20) as client:
            response = await client.post(
                f"https://graph.microsoft.com/v1.0/users/{mailbox}/onenote/sections/{section_id}/pages",
                headers={
                    "Authorization": f"Bearer {token}",
                    "Content-Type": "text/html",
                },
                content=html.encode("utf-8"),
            )
            response.raise_for_status()
            page = response.json()

        links = page.get("links") or {}
        web_url = (links.get("oneNoteWebUrl") or {}).get("href")
        return {
            "id": page.get("id"),
            "title": page.get("title") or title,
            "createdDateTime": page.get("createdDateTime"),
            "lastModifiedDateTime": page.get("lastModifiedDateTime"),
            "webUrl": web_url,
            "section_id": section_id,
            "owner_mailbox": mailbox,
        }

    async def create_email_draft(
        self,
        *,
        agent_name: str,
        to_recipients: list[str],
        subject: str,
        body: str,
    ) -> dict[str, Any]:
        mailbox = self.mailbox_for_agent(agent_name)
        token = await self.get_access_token()
        message = {
            "subject": subject,
            "body": {"contentType": "Text", "content": body},
            "toRecipients": [
                {"emailAddress": {"address": recipient}}
                for recipient in to_recipients
            ],
        }

        async with httpx.AsyncClient(timeout=20) as client:
            response = await client.post(
                f"https://graph.microsoft.com/v1.0/users/{mailbox}/messages",
                headers={"Authorization": f"Bearer {token}"},
                json=message,
            )
            response.raise_for_status()
            return response.json()

    async def send_email(
        self,
        *,
        agent_name: str,
        to_recipients: list[str],
        subject: str,
        body: str,
    ) -> dict[str, Any]:
        if not self.settings.allow_send:
            raise PermissionError("Outbound Microsoft 365 sends are disabled; create drafts for approval instead")

        mailbox = self.mailbox_for_agent(agent_name)
        token = await self.get_access_token()
        payload = {
            "message": {
                "subject": subject,
                "body": {"contentType": "Text", "content": body},
                "toRecipients": [
                    {"emailAddress": {"address": recipient}}
                    for recipient in to_recipients
                ],
            },
            "saveToSentItems": True,
        }

        async with httpx.AsyncClient(timeout=20) as client:
            response = await client.post(
                f"https://graph.microsoft.com/v1.0/users/{mailbox}/sendMail",
                headers={"Authorization": f"Bearer {token}"},
                json=payload,
            )
            response.raise_for_status()
        return {"status": "sent", "agent": agent_name, "mailbox": mailbox}
