# tests/test_meeting_notes_service.py
"""
Tests for the meeting notes publishing service and Nathan's
save_meeting_notes tool. Covers both the happy path (template rendered,
files uploaded, audit recorded) and the failure modes Claude needs to
react to (empty inputs, Graph upload failure).
"""

import asyncio
import unittest
import zipfile
from io import BytesIO
from pathlib import Path
from unittest.mock import AsyncMock, patch

from app.microsoft365 import Microsoft365Settings
from app.services.meeting_notes_service import publish_meeting_notes_to_teams
from app.tools.meeting_notes_tool import save_meeting_notes


# Tests that exercise the Teams-download path force the local-first lookup
# to miss by pointing _ARTIFACTS_ROOT at a path that can't exist.
_NO_LOCAL_ARTIFACTS = Path("__no_local_artifacts_for_tests__")


def _graph_client_with_template():
    """Build an AsyncMock graph client that returns a real DOCX template."""
    # Minimal valid DOCX (a stripped-down version of what's in the
    # existing template tests). Just enough for render_meeting_notes_
    # template_docx to find/replace placeholders.
    docx_bytes = BytesIO()
    with zipfile.ZipFile(docx_bytes, "w") as archive:
        archive.writestr(
            "[Content_Types].xml",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            '<Default Extension="xml" ContentType="application/xml"/>'
            '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
            '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
            "</Types>",
        )
        archive.writestr(
            "_rels/.rels",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
            "</Relationships>",
        )
        archive.writestr(
            "word/document.xml",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:body><w:p><w:r><w:t>{{TITLE}}</w:t></w:r></w:p>'
            '<w:p><w:r><w:t>{{CLIENT_NAME}}</w:t></w:r></w:p>'
            '<w:p><w:r><w:t>{{SUMMARY}}</w:t></w:r></w:p></w:body></w:document>',
        )
    docx_bytes.seek(0)

    graph_client = AsyncMock()
    graph_client.settings = Microsoft365Settings(
        tenant_id="tenant",
        client_id="client",
        client_secret="secret",
        graph_scope="scope",
        webhook_client_state="state",
        allow_send=False,
        agent_mailboxes={"nathan": "nathan@parlayvu.ai"},
    )
    graph_client.download_teams_channel_file.return_value = docx_bytes.getvalue()
    graph_client.upload_teams_channel_file.side_effect = [
        {"id": "md-1", "name": "test.md", "webUrl": "https://sharepoint.example/test.md"},
        {"id": "docx-1", "name": "test.docx", "webUrl": "https://sharepoint.example/test.docx"},
    ]
    return graph_client


class MeetingNotesServiceTests(unittest.TestCase):

    def test_empty_title_raises_value_error(self):
        async def run():
            await publish_meeting_notes_to_teams(
                title="   ",
                summary="A summary",
                client_id="ramair",
            )

        with self.assertRaises(ValueError) as ctx:
            asyncio.run(run())
        self.assertIn("title", str(ctx.exception).lower())

    def test_empty_summary_raises_value_error(self):
        async def run():
            await publish_meeting_notes_to_teams(
                title="A title",
                summary="",
                client_id="ramair",
            )

        with self.assertRaises(ValueError) as ctx:
            asyncio.run(run())
        self.assertIn("summary", str(ctx.exception).lower())

    def test_happy_path_uploads_markdown_and_docx_and_records_audit(self):
        graph_client = _graph_client_with_template()

        with patch(
            "app.services.meeting_notes_service.MicrosoftGraphClient",
            return_value=graph_client,
        ), patch(
            "app.services.meeting_notes_service.memory_get_project_context",
            return_value=None,
        ), patch(
            "app.services.meeting_notes_service.list_clients",
            return_value=[],
        ), patch(
            "app.services.meeting_notes_service.record_generated_output",
            return_value="output-99",
        ) as gen_output, patch(
            "app.services.meeting_notes_service.record_agent_event",
            return_value="event-99",
        ) as event:
            result = asyncio.run(
                publish_meeting_notes_to_teams(
                    title="RamAir Strategy May 24",
                    summary="Discussed Q3 campaign cadence. Approved budget shift to paid social.",
                    client_id="ramair",
                    project_id="ramair-straight-from-the-hart",
                    client_name="RamAir International",
                    channel="tavus_meeting",
                )
            )

        self.assertEqual(result["status"], "published")
        self.assertEqual(result["files"]["markdown"]["id"], "md-1")
        self.assertEqual(result["files"]["docx"]["id"], "docx-1")
        self.assertEqual(result["docx_template"]["status"], "template")
        self.assertEqual(result["memory_output_id"], "output-99")
        self.assertEqual(result["event_id"], "event-99")
        # Audit metadata flows through with the originating channel
        self.assertEqual(event.call_args.kwargs["channel"], "tavus_meeting")
        self.assertEqual(
            gen_output.call_args.kwargs["output_type"],
            "teams_files_meeting_notes",
        )

    def test_falls_back_to_generated_docx_when_template_download_fails(self):
        graph_client = AsyncMock()
        graph_client.settings = Microsoft365Settings(
            tenant_id="tenant",
            client_id="client",
            client_secret="secret",
            graph_scope="scope",
            webhook_client_state="state",
            allow_send=False,
            agent_mailboxes={"nathan": "nathan@parlayvu.ai"},
        )
        graph_client.download_teams_channel_file.side_effect = RuntimeError(
            "template missing"
        )
        graph_client.upload_teams_channel_file.side_effect = [
            {"id": "md-2", "name": "x.md", "webUrl": "https://sharepoint.example/x.md"},
            {"id": "docx-2", "name": "x.docx", "webUrl": "https://sharepoint.example/x.docx"},
        ]

        with patch(
            "app.services.meeting_notes_service.MicrosoftGraphClient",
            return_value=graph_client,
        ), patch(
            "app.services.meeting_notes_service._ARTIFACTS_ROOT",
            _NO_LOCAL_ARTIFACTS,
        ), patch(
            "app.services.meeting_notes_service.record_generated_output",
            return_value="output-1",
        ), patch(
            "app.services.meeting_notes_service.record_agent_event",
            return_value="event-1",
        ):
            result = asyncio.run(
                publish_meeting_notes_to_teams(
                    title="RamAir Strategy",
                    summary="Captured the next steps.",
                    client_id="ramair",
                )
            )

        self.assertEqual(result["status"], "published")
        self.assertEqual(result["docx_template"]["status"], "fallback")
        self.assertIn(
            "template missing",
            result["docx_template"]["fallback_reason"],
        )

    def test_renderer_duplicates_list_paragraphs_and_action_item_rows(self):
        """Pass structured list_items + action_items into the renderer and
        verify the output DOCX has the right number of paragraphs and rows.
        Uses the existing minimal-docx test fixture pattern."""
        from app.microsoft365 import render_meeting_notes_template_docx
        import io
        import zipfile as _zf
        from docx import Document

        # Build a tiny template with: 1 paragraph with {{DECISIONS}}, and
        # a 2-row table where the second row has the action item placeholders.
        template_body = (
            "<w:p><w:r><w:t>Title</w:t></w:r></w:p>"
            "<w:p><w:r><w:t>{{DECISIONS}}</w:t></w:r></w:p>"
            "<w:tbl>"
              "<w:tr><w:tc><w:p><w:r><w:t>OWNER</w:t></w:r></w:p></w:tc>"
                    "<w:tc><w:p><w:r><w:t>ACTION</w:t></w:r></w:p></w:tc>"
                    "<w:tc><w:p><w:r><w:t>DUE</w:t></w:r></w:p></w:tc></w:tr>"
              "<w:tr><w:tc><w:p><w:r><w:t>{{ACTION_OWNER}}</w:t></w:r></w:p></w:tc>"
                    "<w:tc><w:p><w:r><w:t>{{ACTION_ITEM}}</w:t></w:r></w:p></w:tc>"
                    "<w:tc><w:p><w:r><w:t>{{ACTION_DUE}}</w:t></w:r></w:p></w:tc></w:tr>"
            "</w:tbl>"
        )

        # Use the same minimal docx helper format as test_microsoft365
        content_types_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
            '<Default Extension="xml" ContentType="application/xml"/>'
            '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
            "</Types>"
        )
        package_rels_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
            "</Relationships>"
        )
        document_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            f"<w:body>{template_body}<w:sectPr/></w:body>"
            "</w:document>"
        )
        buf = io.BytesIO()
        with _zf.ZipFile(buf, "w") as z:
            z.writestr("[Content_Types].xml", content_types_xml)
            z.writestr("_rels/.rels", package_rels_xml)
            z.writestr("word/document.xml", document_xml)
        template_bytes = buf.getvalue()

        rendered = render_meeting_notes_template_docx(
            template_bytes,
            placeholders={"{{TITLE}}": "ignored"},  # nothing matches; just to exercise the simple-sub path
            list_items={
                "{{DECISIONS}}": ["Decision A.", "Decision B.", "Decision C."],
            },
            action_items=[
                {"owner": "Riley", "action": "File the meeting notes.", "due_date": "EOD today"},
                {"owner": "Dylan", "action": "Stage the site refresh.", "due_date": "May 30"},
            ],
        )

        # Parse the output with python-docx and verify structure
        out_doc = Document(io.BytesIO(rendered))

        decision_paragraphs = [p.text for p in out_doc.paragraphs if "Decision" in p.text]
        self.assertEqual(len(decision_paragraphs), 3, decision_paragraphs)
        self.assertEqual(decision_paragraphs[0], "Decision A.")
        self.assertEqual(decision_paragraphs[1], "Decision B.")
        self.assertEqual(decision_paragraphs[2], "Decision C.")

        # The action items table: header row + 2 data rows
        self.assertEqual(len(out_doc.tables), 1)
        table = out_doc.tables[0]
        self.assertEqual(len(table.rows), 3, "header row plus one row per action item")
        # Data rows have actual values, not placeholders
        row1_text = " ".join(c.text for c in table.rows[1].cells)
        self.assertIn("Riley", row1_text)
        self.assertIn("File the meeting notes.", row1_text)
        self.assertIn("EOD today", row1_text)
        row2_text = " ".join(c.text for c in table.rows[2].cells)
        self.assertIn("Dylan", row2_text)
        self.assertIn("Stage the site refresh.", row2_text)
        self.assertIn("May 30", row2_text)
        # No placeholders survived
        all_text = " ".join(c.text for r in table.rows for c in r.cells)
        self.assertNotIn("{{ACTION_OWNER}}", all_text)
        self.assertNotIn("{{ACTION_ITEM}}", all_text)
        self.assertNotIn("{{ACTION_DUE}}", all_text)

    def test_loads_template_from_local_client_artifacts_first(self):
        """When client_artifacts/<client>/<template> exists locally, the
        service must use that and NOT fall back to Teams download."""
        graph_client = AsyncMock()
        graph_client.settings = Microsoft365Settings(
            tenant_id="tenant",
            client_id="client",
            client_secret="secret",
            graph_scope="scope",
            webhook_client_state="state",
            allow_send=False,
            agent_mailboxes={"nathan": "nathan@parlayvu.ai"},
        )
        # If the service hits Teams, this would explode loudly. We want it
        # to NEVER be called when the local template exists.
        graph_client.download_teams_channel_file.side_effect = AssertionError(
            "Teams download should not be called when local template exists"
        )
        graph_client.upload_teams_channel_file.side_effect = [
            {"id": "md-local", "name": "x.md", "webUrl": "https://sharepoint.example/x.md"},
            {"id": "docx-local", "name": "x.docx", "webUrl": "https://sharepoint.example/x.docx"},
        ]

        # client_id="ramair" exists at client_artifacts/ramair/00_Client_Brief/
        # Templates/RamAir Meeting Notes Template.docx in this repo.
        with patch(
            "app.services.meeting_notes_service.MicrosoftGraphClient",
            return_value=graph_client,
        ), patch(
            "app.services.meeting_notes_service.record_generated_output",
            return_value="output-local",
        ), patch(
            "app.services.meeting_notes_service.record_agent_event",
            return_value="event-local",
        ):
            result = asyncio.run(
                publish_meeting_notes_to_teams(
                    title="Local Template Test",
                    summary="Verifying client_artifacts is the source of truth.",
                    client_id="ramair",
                )
            )

        self.assertEqual(result["status"], "published")
        self.assertEqual(result["docx_template"]["status"], "template")
        # Confirm the template came from local artifacts, not Teams
        self.assertEqual(result["docx_template"]["source"], "client_artifacts")
        self.assertIn("client_artifacts/ramair", result["docx_template"]["path"])
        graph_client.download_teams_channel_file.assert_not_awaited()


class SaveMeetingNotesToolTests(unittest.TestCase):
    """The Tavus-facing tool wraps the service and translates failures
    into Nathan-friendly result dicts. Claude needs structured errors,
    not exceptions, to stay graceful mid-meeting."""

    def test_happy_path_returns_saved_status_with_urls_and_message(self):
        fake_service_result = {
            "status": "published",
            "files": {
                "markdown": {"webUrl": "https://sharepoint.example/notes.md"},
                "docx": {"webUrl": "https://sharepoint.example/notes.docx"},
            },
            "memory_output_id": "out-1",
            "event_id": "evt-1",
        }

        with patch(
            "app.tools.meeting_notes_tool.publish_meeting_notes_to_teams",
            new=AsyncMock(return_value=fake_service_result),
        ) as publish_mock:
            result = asyncio.run(
                save_meeting_notes(
                    title="RamAir Wrap",
                    summary="Discussed cadence.",
                    client_id="ramair",
                )
            )

        # Service was called with the tavus_meeting channel and nathan agent
        self.assertEqual(publish_mock.call_args.kwargs["channel"], "tavus_meeting")
        self.assertEqual(publish_mock.call_args.kwargs["agent_name"], "nathan")
        self.assertEqual(result["status"], "saved")
        self.assertEqual(result["title"], "RamAir Wrap")
        self.assertEqual(result["client_id"], "ramair")
        self.assertEqual(result["markdown_url"], "https://sharepoint.example/notes.md")
        self.assertEqual(result["docx_url"], "https://sharepoint.example/notes.docx")
        self.assertEqual(result["memory_output_id"], "out-1")
        self.assertIn("ramair", result["message"].lower())

    def test_value_error_returns_failed_status_not_exception(self):
        with patch(
            "app.tools.meeting_notes_tool.publish_meeting_notes_to_teams",
            new=AsyncMock(side_effect=ValueError("Meeting note title is required")),
        ):
            result = asyncio.run(
                save_meeting_notes(
                    title="",
                    summary="A summary",
                    client_id="ramair",
                )
            )

        self.assertEqual(result["status"], "failed")
        self.assertIn("title", result["error"].lower())
        # Friendly message Nathan can read aloud, not a stack trace
        self.assertIn("can't save", result["message"].lower())

    def test_graph_failure_returns_failed_status_with_followup_message(self):
        with patch(
            "app.tools.meeting_notes_tool.publish_meeting_notes_to_teams",
            new=AsyncMock(side_effect=RuntimeError("Graph 503")),
        ):
            result = asyncio.run(
                save_meeting_notes(
                    title="RamAir Wrap",
                    summary="Discussed cadence.",
                    client_id="ramair",
                )
            )

        self.assertEqual(result["status"], "failed")
        self.assertIn("Graph 503", result["error"])
        # The "I'll have someone follow up" line is what Nathan reads on failure
        self.assertIn("follow up", result["message"].lower())


if __name__ == "__main__":
    unittest.main()
