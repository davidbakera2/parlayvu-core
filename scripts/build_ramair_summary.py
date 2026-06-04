"""Build the RamAir / CIRI Journal Marketing Strategy meeting summary docx."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_PATH = (
    r"C:\Users\DavidBaker\OneDrive - Baker Strategy Group\Documents"
    r"\RamAir Meeting Summary - 2026-05-27.docx"
)

doc = Document()

# Base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)


def add_subtitle(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)


def add_para(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(text, *, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(2)


# ---------- Header ----------
add_title("Meeting Summary — RamAir Marketing Strategy")
add_subtitle("CIRI Journal Marketing Strategy | May 27, 2026 | Duration: 1h 27m")

# ---------- Attendees ----------
add_h2("Attendees")
attendees = [
    ("Sally Compton", "Marketing lead, RamAir International (also operates Think Wild nonprofit, Bend, OR)"),
    ("Andrew Luckey", "RamAir International — co-author of the CIRI white paper"),
    ("Cole Younger", "RamAir International — sales"),
    ("David Baker", "ParlayVU — marketing strategist (external)"),
    ("Nathan Ellis", "Note-taker / transcription"),
]
for name, role in attendees:
    add_bullet(f" — {role}", bold_lead=name)

# ---------- Executive Summary ----------
add_h2("Executive Summary")
add_para(
    "The team convened to map a go-to-market strategy off the back of RamAir’s newly published, "
    "peer-reviewed CIRI Journal research paper validating positive air duct cleaning. Sally framed the "
    "paper as a watershed moment — “our flag on the moon” — after 18 years of pioneering the "
    "category against the dominant negative-air orthodoxy. With the science now anchored by Element "
    "Labs testing and dual peer review through CIRI/IICRC channels, the conversation focused on how to "
    "convert that credibility into awareness, leads, and ultimately equipment sales."
)
add_para(
    "Alignment was reached on rebranding the equipment as the “Ram Air Positive Air Duct Cleaning "
    "System” to own the trademarked term and differentiate from copycats (Airwave / Hydro Force, "
    "Duct Pro / Pro Air). On tone, David and Sally agreed to acknowledge negative air duct cleaning but "
    "not dwell on it — address it, dismiss it, move on — while leaning into a stronger, headline-ready "
    "claim (“new study proves [this method] can damage your HVAC system”) tailored per audience: "
    "homeowners/facility managers, contractors/restorers, and the trade industry itself."
)
add_para(
    "The launch plan starts organic: a two-track press release — local (Bend / Guarantee-branded) "
    "and broader RamAir (Pacific Northwest, then California, then national) — distributed through "
    "Sally’s existing media contacts and trade pubs (R&R, Clean Facts, Cleaning & Restoration, "
    "Snips), with a media kit anchored by a dedicated video interview between David and John Miles "
    "(confirmed Friday, 1pm ET). Paid PR outreach and a Wikipedia page were sequenced for after organic "
    "coverage and backlinks accrue."
)
add_para(
    "LinkedIn was deliberately deferred — not abandoned — to avoid surfacing naysayer pushback in "
    "front of IICRC, whose buy-in Sally is cultivating quietly through her relationship with Holly. In "
    "parallel, the team committed to restarting a multi-channel content engine: daily social clips "
    "(TikTok, Instagram, YouTube, Twitter) built from a 14-video rotation, sourced from a shared "
    "iPhone album that field techs (Dakota, Spencer) feed daily; an email newsletter restart to the "
    "full RamAir customer list; an HVAC-contractor email campaign pitching duct cleaning as a high-margin "
    "add-on; and a crowdsourced testimonial compilation. Cole emphasized that all messaging should "
    "ladder back to the three buyer hooks: low startup cost, high earning potential, and ease of use."
)

# ---------- Action Items ----------
add_h2("Action Items")

action_groups = [
    ("Brand & Positioning", [
        ("Sally", "Standardize “Ram Air Positive Air Duct Cleaning System” across collateral, website, and sales tools."),
        ("Sally / David", "Define audience-specific headlines (homeowner, contractor, trade) before press release goes out; anchor consumer message around “new study proves [X] can damage your HVAC system.”"),
        ("David / Sally", "Distill four key messages from the CIRI paper to drive the John Miles interview and downstream marketing talking points."),
    ]),
    ("Press Release & PR", [
        ("Sally", "Draft two press releases: (1) local / Guarantee-focused for Central Oregon outlets (CO Daily, KTVZ, etc.); (2) broader RamAir release for Pacific Northwest → California → national."),
        ("Sally", "Distribute to existing organic media contact lists; specifically loop in Kayla at R&R Magazine, plus Clean Facts, Cleaning & Restoration, and Snips."),
        ("Sally", "Solicit endorsement quotes from peer reviewers (Dr. Carrie LeSage, Dr. Brandon White) and reference contacts (Brad Smith, Clay Fernandez, others) to amplify the release through their networks."),
        ("Sally / David", "Defer paid PR-firm engagement until organic results have been measured."),
    ]),
    ("John Miles Interview — Friday", [
        ("Sally", "Email John Miles and David Baker to confirm Friday recording at 1pm ET / 10am PT; copy David."),
        ("David", "Send John the calendar invite and meeting link once Sally’s confirmation email lands."),
        ("David / Sally", "Produce a dedicated paper-focused interview as a Straight From the Heart episode; capture B-roll/clips suitable for inclusion in the press-release media kit."),
    ]),
    ("LinkedIn & ICRC", [
        ("Sally / Andrew", "Track down Fred (or Paul) to grant David Baker admin access to the existing RamAir International LinkedIn page — tee it up even though posting is paused."),
        ("Sally", "Continue the soft-touch ICRC strategy via Holly; do not initiate LinkedIn posts about the paper until the ICRC dynamic is clearer."),
        ("Sally / David", "Revisit the LinkedIn decision after initial PR coverage lands — door is open, not closed."),
    ]),
    ("Social Content Engine", [
        ("Sally / Cole", "Build a 14-video rotation of short-form clips (system-in-action visuals + text overlays) and post daily across TikTok, Instagram, YouTube Shorts, and Twitter; cycle on a 14-day loop."),
        ("Sally", "Lead messaging on the three buyer hooks: low startup cost, high earning potential, ease of use — plus repurpose viral 2M / 10M-view clips into paid ads."),
        ("Sally / Cole", "Stand up a shared iPhone album; require Dakota and Spencer to capture one job-site video per day as part of their job duties; explore a per-sale bonus tied to footage that converts."),
        ("Sally", "Visit job sites ~monthly to capture third-party video content directly."),
    ]),
    ("Email & Customer Reactivation", [
        ("Sally / Cole", "Restart the monthly RamAir customer newsletter — feature the CIRI paper, latest podcast episode, monthly blog, and a recurring spotlight."),
        ("Sally / Cole", "Compose an email request for short customer testimonial clips, paired with a small incentive (e.g., branded swag); assemble into a compilation video."),
        ("Cole / David", "Locate the dormant HVAC-contractor email list; launch an outreach campaign pitching duct cleaning as a high-margin add-on service, modeled on the two Bend HVAC partners already converting."),
    ]),
    ("Website & Podcast", [
        ("Cole", "Confirm chatbox status on the RamAir and Guarantee sites; reinstate Talk2 (Parrot) live-chat with mobile notifications routed to Cole (RamAir) and Sally (Guarantee)."),
        ("David", "Create differentiated YouTube thumbnails per Straight From the Heart episode — consistent branding, but episode number, title, and guest photo unique to each."),
        ("David / Cole", "Continue lining up podcast guests; discuss future guest cadence and inbound-request strategy offline."),
    ]),
    ("Wikipedia (sequenced)", [
        ("Sally", "Begin the RamAir Wikipedia page build-out only after PR coverage has generated external backlinks and third-party references."),
    ]),
]

for group_title, items in action_groups:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(group_title)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    for owner, task in items:
        add_bullet(f" — {task}", bold_lead=owner)

doc.save(OUT_PATH)
print(f"Wrote {OUT_PATH}")
